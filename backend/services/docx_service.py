"""Structured DOCX resume export mirroring original section order."""
import io
from typing import Optional

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    Document = None

from models.schemas import CandidateProfile, ResumePreviewResponse, TailoredExperienceEntry
from services.guardrails.constants import MAX_EXPERIENCE_BULLETS
from services.pdf_service import DEFAULT_SECTION_ORDER


def _hex_to_rgb(hex_color: str) -> tuple[int, int, int]:
    h = hex_color.lstrip("#")
    if len(h) == 3:
        h = "".join(c * 2 for c in h)
    try:
        return int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
    except (ValueError, IndexError):
        return 16, 185, 129


def _add_section_heading(doc: "Document", title: str, accent_hex: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(11)
    r, g, b = _hex_to_rgb(accent_hex)
    run.font.color.rgb = RGBColor(r, g, b)


def generate_docx_from_package(
    profile: CandidateProfile,
    package: ResumePreviewResponse,
    *,
    accent_hex: str = "#10b981",
) -> bytes:
    if Document is None:
        raise RuntimeError("python-docx is not installed.")

    order = [s for s in (package.section_order or DEFAULT_SECTION_ORDER) if s in DEFAULT_SECTION_ORDER]
    for sec in DEFAULT_SECTION_ORDER:
        if sec not in order:
            order.append(sec)

    highlighted = {p.lower() for p in (package.highlighted_projects or [])}
    doc = Document()

    name_p = doc.add_paragraph()
    name_run = name_p.add_run(profile.name or "Candidate")
    name_run.bold = True
    name_run.font.size = Pt(22)

    contact_parts = [x for x in [profile.email, profile.phone, profile.location] if x]
    if contact_parts:
        cp = doc.add_paragraph("  |  ".join(contact_parts))
        cp.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        for run in cp.runs:
            run.font.size = Pt(9)

    def summary_block() -> None:
        if not package.tailored_summary:
            return
        _add_section_heading(doc, "Professional Summary", accent_hex)
        p = doc.add_paragraph(package.tailored_summary)
        for run in p.runs:
            run.font.size = Pt(10)

    def skills_block() -> None:
        if not package.ordered_skills:
            return
        _add_section_heading(doc, "Skills", accent_hex)
        chunks = [package.ordered_skills[i : i + 5] for i in range(0, len(package.ordered_skills), 5)]
        for chunk in chunks:
            p = doc.add_paragraph(", ".join(chunk), style="List Bullet")
            for run in p.runs:
                run.font.size = Pt(10)

    def experience_block() -> None:
        entries = package.tailored_experience
        if not entries and profile.experience:
            entries = [
                TailoredExperienceEntry(
                    company=exp.company,
                    role=exp.role,
                    duration=exp.duration,
                    bullets=exp.description[:MAX_EXPERIENCE_BULLETS],
                )
                for exp in profile.experience
            ]
        if not entries:
            return
        _add_section_heading(doc, "Experience", accent_hex)
        for exp in entries:
            title_p = doc.add_paragraph()
            role_run = title_p.add_run(exp.role)
            role_run.bold = True
            role_run.font.size = Pt(10)
            if exp.duration:
                dur_run = title_p.add_run(f"  —  {exp.duration}")
                dur_run.font.size = Pt(9)
            comp_p = doc.add_paragraph(exp.company)
            for run in comp_p.runs:
                run.italic = True
                run.font.size = Pt(9)
            for bullet in exp.bullets[:MAX_EXPERIENCE_BULLETS]:
                bp = doc.add_paragraph(bullet, style="List Bullet")
                for run in bp.runs:
                    run.font.size = Pt(10)

    def projects_block() -> None:
        if not profile.projects:
            return
        _add_section_heading(doc, "Projects", accent_hex)
        ordered = sorted(profile.projects, key=lambda p: 0 if p.name.lower() in highlighted else 1)
        for proj in ordered[:8]:
            tech = ", ".join(proj.technologies) if proj.technologies else ""
            title = proj.name + (f"  [{tech}]" if tech else "")
            tp = doc.add_paragraph()
            tr = tp.add_run(title)
            tr.bold = True
            tr.font.size = Pt(10)
            if proj.description:
                dp = doc.add_paragraph(proj.description)
                for run in dp.runs:
                    run.font.size = Pt(10)

    def education_block() -> None:
        if not profile.education:
            return
        _add_section_heading(doc, "Education", accent_hex)
        for edu in profile.education:
            tp = doc.add_paragraph()
            dr = tp.add_run(edu.degree)
            dr.bold = True
            dr.font.size = Pt(10)
            if edu.year:
                yr = tp.add_run(f"  —  {edu.year}")
                yr.font.size = Pt(9)
            ip = doc.add_paragraph(edu.institution)
            for run in ip.runs:
                run.italic = True
                run.font.size = Pt(9)

    renderers = {
        "summary": summary_block,
        "skills": skills_block,
        "experience": experience_block,
        "projects": projects_block,
        "education": education_block,
    }
    for sec in order:
        renderers[sec]()

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
